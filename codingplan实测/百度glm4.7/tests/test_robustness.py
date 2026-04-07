# -*- coding: utf-8 -*-
"""
水印鲁棒性测试
"""

import os
import sys
import tempfile
import shutil
from pathlib import Path
import unittest

# 添加项目根目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.core.watermark import DocxWatermarkTool


class RobustnessTestBase(unittest.TestCase):
    """鲁棒性测试基类"""

    @classmethod
    def setUpClass(cls):
        """创建测试文档"""
        cls.temp_dir = tempfile.mkdtemp()
        cls.test_docx = cls._create_test_docx()
        cls.tool = DocxWatermarkTool('robustness_test')
        cls.watermarked_docx = os.path.join(cls.temp_dir, 'watermarked.docx')

        # 预先嵌入水印
        cls.tool.embed_watermark(
            cls.test_docx,
            cls.watermarked_docx,
            '鲁棒性测试用户',
            '测试部门',
            '测试项目'
        )

    @classmethod
    def tearDownClass(cls):
        """清理测试文件"""
        shutil.rmtree(cls.temp_dir, ignore_errors=True)

    @classmethod
    def _create_test_docx(cls):
        """创建测试用的 DOCX 文件"""
        from docx import Document

        doc = Document()
        doc.add_heading('鲁棒性测试文档', level=1)

        for i in range(20):
            doc.add_paragraph(f'段落 {i+1}：这是测试内容，用于验证水印的鲁棒性。中文和English混合。')

        test_path = os.path.join(cls.temp_dir, 'robustness_test.docx')
        doc.save(test_path)
        return test_path

    def _verify_watermark(self, file_path):
        """验证水印"""
        result = self.tool.analyze_docx(file_path)
        return result['success'] and result['has_watermark']


class TestPartialDeletion(RobustnessTestBase):
    """部分删除测试"""

    def test_delete_first_paragraph(self):
        """测试删除第一个段落后水印保留"""
        from docx import Document
        import zipfile

        # 复制文件
        modified_path = os.path.join(self.temp_dir, 'deleted_first.docx')
        shutil.copy(self.watermarked_docx, modified_path)

        # 删除第一个段落
        doc = Document(modified_path)
        if doc.paragraphs:
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)
        doc.save(modified_path)

        # 验证水印
        self.assertTrue(self._verify_watermark(modified_path))

    def test_delete_half_paragraphs(self):
        """测试删除一半段落后水印保留"""
        from docx import Document

        modified_path = os.path.join(self.temp_dir, 'deleted_half.docx')
        shutil.copy(self.watermarked_docx, modified_path)

        doc = Document(modified_path)
        total = len(doc.paragraphs)
        delete_count = total // 2

        # 从前向后删除
        for i in range(delete_count):
            if doc.paragraphs:
                p = doc.paragraphs[0]._element
                p.getparent().remove(p)

        doc.save(modified_path)

        # 验证水印
        self.assertTrue(self._verify_watermark(modified_path))


class TestFormatChanges(RobustnessTestBase):
    """格式修改测试"""

    def test_change_font(self):
        """测试修改字体后水印保留"""
        from docx import Document
        from docx.shared import Pt

        modified_path = os.path.join(self.temp_dir, 'changed_font.docx')
        shutil.copy(self.watermarked_docx, modified_path)

        doc = Document(modified_path)

        for para in doc.paragraphs:
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(14)

        doc.save(modified_path)

        self.assertTrue(self._verify_watermark(modified_path))

    def test_add_formatting(self):
        """测试添加格式后水印保留"""
        from docx import Document
        from docx.shared import Pt, RGBColor

        modified_path = os.path.join(self.temp_dir, 'added_format.docx')
        shutil.copy(self.watermarked_docx, modified_path)

        doc = Document(modified_path)

        for para in doc.paragraphs:
            for run in para.runs:
                run.bold = True
                run.italic = True
                run.font.color.rgb = RGBColor(255, 0, 0)

        doc.save(modified_path)

        self.assertTrue(self._verify_watermark(modified_path))


class TestBackupLayers(RobustnessTestBase):
    """备份层测试"""

    def test_custom_xml_backup(self):
        """测试 custom.xml 备份层"""
        import zipfile

        # 检查备份层是否存在
        with zipfile.ZipFile(self.watermarked_docx, 'r') as z:
            self.assertIn('docProps/custom.xml', z.namelist())

        # 验证能从备份层提取
        result = self.tool.analyze_docx(self.watermarked_docx)
        self.assertTrue(result['success'])

    def test_settings_xml_backup(self):
        """测试 settings.xml 备份层"""
        import zipfile

        # 检查备份层是否存在
        with zipfile.ZipFile(self.watermarked_docx, 'r') as z:
            content = z.read('word/settings.xml').decode('utf-8')
            self.assertIn('wm_store:', content)

    def test_backup_xml_file(self):
        """测试独立备份文件"""
        import zipfile

        # 检查备份文件是否存在
        with zipfile.ZipFile(self.watermarked_docx, 'r') as z:
            self.assertIn('word/watermark_backup.xml', z.namelist())


class TestSimulatedWPSRewrite(RobustnessTestBase):
    """模拟 WPS 重写测试"""

    def test_zero_width_removed(self):
        """测试零宽字符被清除后的恢复"""
        import zipfile
        import re

        # 创建模拟 WPS 清除零宽字符的版本
        modified_path = os.path.join(self.temp_dir, 'zw_removed.docx')

        with zipfile.ZipFile(self.watermarked_docx, 'r') as zin:
            with zipfile.ZipFile(modified_path, 'w') as zout:
                for item in zin.namelist():
                    content = zin.read(item)

                    if item == 'word/document.xml':
                        # 清除所有零宽字符
                        text = content.decode('utf-8')
                        text = re.sub(r'[\u200b\u200c]', '', text)
                        content = text.encode('utf-8')

                    zout.writestr(item, content)

        # 验证能从备份层恢复
        result = self.tool.analyze_docx(modified_path)
        self.assertTrue(result['success'])
        self.assertIn(result['source'], ['custom.xml', 'Content_Types.xml',
                                         'settings.xml', 'header.xml',
                                         'watermark_backup.xml'])


if __name__ == '__main__':
    unittest.main()
