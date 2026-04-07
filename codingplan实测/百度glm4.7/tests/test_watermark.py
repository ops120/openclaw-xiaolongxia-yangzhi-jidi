# -*- coding: utf-8 -*-
"""
水印功能测试
"""

import os
import sys
import tempfile
import shutil
from pathlib import Path
import unittest

# 添加项目根目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.core.watermark import DocxWatermarkTool
from src.core.crypto import CryptoManager


class TestCryptoManager(unittest.TestCase):
    """加密管理器测试"""

    def test_encrypt_decrypt(self):
        """测试加密解密"""
        crypto = CryptoManager('test_password')
        original = 'Hello, 世界!'

        encrypted = crypto.encrypt(original)
        decrypted = crypto.decrypt(encrypted)

        self.assertEqual(original, decrypted)

    def test_base64_encrypt_decrypt(self):
        """测试 base64 加密解密"""
        crypto = CryptoManager('test_password')
        original = '测试数据 123'

        encrypted_b64 = crypto.encrypt_to_base64(original)
        decrypted = crypto.decrypt_from_base64(encrypted_b64)

        self.assertEqual(original, decrypted)

    def test_default_password(self):
        """测试默认密码"""
        crypto1 = CryptoManager()
        crypto2 = CryptoManager()

        # 使用相同的默认密码，应该能互相解密
        encrypted = crypto1.encrypt('test')
        decrypted = crypto2.decrypt(encrypted)

        self.assertEqual('test', decrypted)

    def test_password_mismatch(self):
        """测试密码不匹配"""
        crypto1 = CryptoManager('password1')
        crypto2 = CryptoManager('password2')

        encrypted = crypto1.encrypt('test')

        with self.assertRaises(Exception):
            crypto2.decrypt(encrypted)


class TestWatermarkTool(unittest.TestCase):
    """水印工具测试"""

    @classmethod
    def setUpClass(cls):
        """创建测试文档"""
        cls.temp_dir = tempfile.mkdtemp()
        cls.test_docx = cls._create_test_docx()

    @classmethod
    def tearDownClass(cls):
        """清理测试文件"""
        shutil.rmtree(cls.temp_dir, ignore_errors=True)

    @classmethod
    def _create_test_docx(cls):
        """创建测试用的 DOCX 文件"""
        from docx import Document

        doc = Document()
        doc.add_heading('测试文档', level=1)
        doc.add_paragraph('这是第一个段落，用于测试水印嵌入功能。')
        doc.add_paragraph('这是第二个段落，包含一些中文和 English 内容。')
        doc.add_paragraph('这是第三个段落，测试水印的鲁棒性。')

        for i in range(10):
            doc.add_paragraph(f'段落 {i+4}：测试内容填充，确保有足够的段落进行水印嵌入。')

        test_path = os.path.join(cls.temp_dir, 'test_document.docx')
        doc.save(test_path)
        return test_path

    def test_embed_and_extract(self):
        """测试水印嵌入和提取"""
        tool = DocxWatermarkTool('test_password')

        output_path = os.path.join(self.temp_dir, 'watermarked.docx')

        # 嵌入水印
        result = tool.embed_watermark(
            self.test_docx,
            output_path,
            '张三-测试',
            '测试部门',
            '测试项目'
        )

        self.assertTrue(result['success'])
        self.assertGreater(result['paragraphs_processed'], 0)
        self.assertTrue(os.path.exists(output_path))

        # 提取水印
        analyze_result = tool.analyze_docx(output_path)

        self.assertTrue(analyze_result['success'])
        self.assertTrue(analyze_result['has_watermark'])
        self.assertEqual(analyze_result['watermark_data']['uid'], '张三-测试')
        self.assertEqual(analyze_result['watermark_data']['department'], '测试部门')

    def test_extract_without_password(self):
        """测试无密码提取"""
        # 使用默认密码嵌入
        tool1 = DocxWatermarkTool()
        output_path = os.path.join(self.temp_dir, 'default_password.docx')

        tool1.embed_watermark(
            self.test_docx,
            output_path,
            '默认密码测试'
        )

        # 使用默认密码提取
        tool2 = DocxWatermarkTool()
        result = tool2.analyze_docx(output_path)

        self.assertTrue(result['success'])
        self.assertEqual(result['watermark_data']['uid'], '默认密码测试')

    def test_backup_layers(self):
        """测试备份层"""
        tool = DocxWatermarkTool('backup_test')
        output_path = os.path.join(self.temp_dir, 'backup_test.docx')

        result = tool.embed_watermark(
            self.test_docx,
            output_path,
            '备份层测试',
            '测试部'
        )

        self.assertTrue(result['success'])
        # 应该至少有3个备份层
        self.assertGreaterEqual(len(result['backup_layers']), 3)

    def test_key_mismatch(self):
        """测试密钥不匹配"""
        tool1 = DocxWatermarkTool('password1')
        output_path = os.path.join(self.temp_dir, 'mismatch_test.docx')

        tool1.embed_watermark(
            self.test_docx,
            output_path,
            '密钥不匹配测试'
        )

        # 使用不同密钥提取
        tool2 = DocxWatermarkTool('password2')
        result = tool2.analyze_docx(output_path)

        # 应该失败或返回错误
        self.assertFalse(result['success'] and result['has_watermark'])


class TestDatabase(unittest.TestCase):
    """数据库测试"""

    def setUp(self):
        """创建临时数据库"""
        import tempfile
        self.temp_db = tempfile.NamedTemporaryFile(suffix='.db', delete=False)
        self.temp_db.close()

    def tearDown(self):
        """删除临时数据库"""
        os.unlink(self.temp_db.name)

    def test_key_crud(self):
        """测试密钥 CRUD 操作"""
        from src.db.models import Database, KeyManager

        db = Database(self.temp_db.name)
        manager = KeyManager(db)

        # 创建
        key = manager.create_key('test_key', 'test_password')
        self.assertEqual(key['key_name'], 'test_key')

        # 读取
        retrieved = manager.get_key('test_key')
        self.assertEqual(retrieved['password'], 'test_password')

        # 获取列表
        keys = manager.get_all_keys()
        self.assertEqual(len(keys), 1)

        # 删除
        self.assertTrue(manager.delete_key('test_key'))
        keys = manager.get_all_keys()
        self.assertEqual(len(keys), 0)

    def test_key_export_import(self):
        """测试密钥导出导入"""
        from src.db.models import Database, KeyManager
        import json

        db = Database(self.temp_db.name)
        manager = KeyManager(db)

        # 创建并导出
        manager.create_key('export_test', 'export_password')
        exported = manager.export_key('export_test')

        self.assertIsNotNone(exported)
        data = json.loads(exported)
        self.assertEqual(data['key_name'], 'export_test')

        # 删除后重新导入
        manager.delete_key('export_test')
        manager.import_key(exported)

        retrieved = manager.get_key('export_test')
        self.assertEqual(retrieved['password'], 'export_password')


if __name__ == '__main__':
    unittest.main()
